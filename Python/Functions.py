import pandas as pd
import numpy as np
import re
import requests
import pyreadstat
import statsmodels
import warnings
from itertools import combinations
from itertools import product
from scipy import stats
from tqdm import tqdm
from scipy.stats import chi2_contingency
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from openpyxl.utils.dataframe import dataframe_to_rows
from statsmodels.stats.weightstats import ttest_ind
warnings.filterwarnings("ignore")

def analysis(file):

   values, codebook = pyreadstat.read_sav(file, apply_value_formats = False)
   labels = pyreadstat.read_sav(file, apply_value_formats = True)[0]

   for combination in (list(combinations(list(product(values["Group"].unique(), values["Time"].unique())), 2))): #tqdm
        
        class sample:
            one = subsample(combination[0][0], combination[0][1], "weight_a", values, labels) # Add multiweight support
            two = subsample(combination[1][0], combination[1][1], "weight_a", values, labels)
            name = comparison_name(one, two)
            metadata = codebook ##.assign(Overall = np.nan)
            paired = combination[0][0] == combination[1][0]
        
        sample.one.values.set_index(sample.one.values["IDs"], inplace=True)
        sample.two.values.set_index(sample.two.values["IDs"], inplace=True)
        sample.one.labels.set_index(sample.one.values["IDs"], inplace=True)
        sample.two.labels.set_index(sample.two.values["IDs"], inplace=True)
        
        analysis_tables(sample, "Nominal")
        analysis_tables(sample, "Ordinal")

   print("Analysis complete.")

def analysis_tables(sample, type):
    
    sample.crosstabs = pd.DataFrame()
    
    sample.metadata.variable_measure.pop('Group', None)
    sample.metadata.variable_measure.pop('Time', None)
    nominal_variables = [key for key, measure in sample.metadata.variable_measure.items() if measure == 'nominal']
    
    if type == "Nominal":
        ordinal_variables = [1]
    if type == "Ordinal":
        ordinal_variables = [key for key, measure in sample.metadata.variable_measure.items() if measure == 'ordinal']
    
    for nominal_variable in nominal_variables:
        for ordinal_variable in reversed(ordinal_variables):

            if type == "Nominal":
                sample.crosstab = nominal_crosstab(sample, nominal_variable)
            if type == "Ordinal":
                sample.crosstab = ordinal_crosstab(sample, nominal_variable, ordinal_variable)
            
            sample.crosstab.loc[-1] = [pd.NA] * len(sample.crosstab.columns)
            sample.crosstab.index += 1
            sample.crosstab.sort_index(inplace = True)
            
            sample.crosstabs = combine_crosstabs(sample.crosstab, sample.crosstabs)
        
        if type == "Ordinal":
            
            sample.crosstabs = sample.crosstabs.iloc[1:]
            second_header = [""] * len(sample.crosstabs.columns)
            second_header[2] = sample.one.name
            second_header[int((len(sample.crosstabs.columns) - 2) / 3) + 2] = sample.two.name
            second_header[int((len(sample.crosstabs.columns) - 2) / 3) * 2 + 2] = "Difference"
            multi_col = pd.MultiIndex.from_tuples([(second_header[i], col) for i, col in enumerate(sample.crosstabs.columns)])
            sample.crosstabs.columns = multi_col
            
            write_excel(sample, document_title(sample, "Tables", type, sample.metadata.column_labels[sample.metadata.column_names.index(nominal_variable)]))
            #write_word(sample, document_title(sample, "Tables", type, sample.metadata.column_labels[sample.metadata.column_names.index(nominal_variable)]))
            #ordinal_report(sample)
            sample.crosstabs = pd.DataFrame()
        
    if type == "Nominal":

        write_excel(sample, document_title(sample, "Tables", type, nominal_variable))
        #write_word(sample, document_title(sample, "Tables", type, nominal_variable))

def nominal_crosstab(sample, nominal_variable):
    
    sample.one.crosstab = create_crosstab(
        type = "Nominal",
        data = sample.one.labels,
        index = nominal_variable,
        columns = "Overall",
        weight = sample.one.weight)
    
    sample.two.crosstab = create_crosstab(
        type = "Nominal",
        data = sample.two.labels,
        index = nominal_variable,
        columns = "Overall",
        weight = sample.two.weight)
    
    sample.crosstab = pd.concat([sample.one.crosstab, sample.two.crosstab], axis=1)
    
    sample.crosstab = sample.crosstab.reset_index()
    sample.crosstab.insert(0, 'Category', np.nan)
    sample.crosstab.columns = ["Category",
                        "Group",
                        add_sample_size(sample.one.name, sample.one.values[nominal_variable]),
                        add_sample_size(sample.two.name, sample.two.values[nominal_variable])]

    sample.crosstab.loc[0, 'Category'] = test_chi(
        variable = sample.metadata.column_labels[sample.metadata.column_names.index(nominal_variable)],
        observed = pd.crosstab(index = sample.one.labels[nominal_variable], columns = 1, values = sample.one.labels[sample.one.weight], aggfunc = 'sum'),
        expected = pd.crosstab(index = sample.two.labels[nominal_variable], columns = 1, values = sample.two.labels[sample.one.weight], aggfunc = 'sum'))

    return sample.crosstab

def ordinal_crosstab(sample, nominal_variable, ordinal_variable):
    
    sample.one.crosstab = create_crosstab(
        type = "Ordinal",
        data = sample.one.labels,
        index = ordinal_variable,
        columns = nominal_variable,
        weight = sample.one.weight)
    
    sample.two.crosstab = create_crosstab(
        type = "Ordinal",
        data = sample.two.labels,
        index = ordinal_variable,
        columns = nominal_variable,
        weight = sample.two.weight)
    
    sample.crosstab = pd.concat([sample.one.crosstab, sample.two.crosstab, sample.two.crosstab - sample.one.crosstab], axis=1)
    sample.crosstab = sample.crosstab.round(1).apply(lambda x: x).applymap(lambda x: f"{x}%")
    
    sample.crosstab = add_crosstab_tests(sample, nominal_variable, ordinal_variable)
    
    sample.crosstab = sample.crosstab.reset_index()
    sample.crosstab.insert(0, 'Variable', np.nan)
    crosstab_header = sample.crosstab.columns.tolist()
    crosstab_header[0] = "Variable"
    crosstab_header[1] = "Prompt and Responses"
    sample.crosstab.columns = crosstab_header
    sample.crosstab.iloc[0,0] = ordinal_variable
    sample.crosstab.iloc[0,1] = sample.metadata.column_labels[sample.metadata.column_names.index(ordinal_variable)]
    
    return sample.crosstab

def ordinal_report(Crosstabs, Outputs, File_Name, Name_Group, Template, Document_Title, Type, Demographic_Category, API_Key, Group1):

    def return_text(change):
        if change < 0:
            return "decreased"
        elif change > 0:
            return "increased"
        elif change == 0:
            return "did not change"

    def stance(support_percent, opposition_percent, stance_negative, stance_positive):
        if support_percent >= 0.67:
            return f"a supermajority selected {stance_positive.lower()}"
        elif support_percent > 0.5:
            return f"a majority selected {stance_positive.lower()}"
        elif opposition_percent > 0.67:
            return f"a supermajority selected {stance_negative.lower()}"
        elif opposition_percent >= 0.5:
            return f"a majority selected {stance_negative.lower()}"
        else:
            return "there was no majority"

    def stance_percentage(support_percent, opposition_percent, stance_negative, stance_positive):
        if support_percent >= 0.67:
            return f"a supermajority selected {stance_positive.lower()} ({format_percentage(support_percent)})"
        elif support_percent > 0.5:
            return f"a majority selected {stance_positive.lower()} ({format_percentage(support_percent)})"
        elif opposition_percent > 0.67:
            return f"a supermajority selected {stance_negative.lower()} ({format_percentage(opposition_percent)})"
        elif opposition_percent >= 0.5:
            return f"a majority selected {stance_negative.lower()} ({format_percentage(opposition_percent)})"
        else:
            return f"there was no majority for {stance_positive.lower()} ({format_percentage(support_percent)}) or {stance_negative.lower()} ({format_percentage(opposition_percent)})"

    def convert_percentage(percentage):
        return float(percentage.strip('%')) / 100

    def format_percentage(percent):
        return f"{abs(percent * 100)}%"

    def add_p_in_parenthesis(s):
        if '(' in s:
            return re.sub(r'\(([^)]+)\)', r'(P = \1)', s)
        else:
            return s

    def extract_first_numeric(s):
        parts = re.split(r' |\(', s)
        return float(parts[0])

    legend = Crosstabs[2:(Crosstabs.index("Prompt and Responses")-6)][1]
    sample_sizes = Crosstabs[Crosstabs.index("Prompt and Responses")-2][2:]

    if 2 <= len(legend) < 7:
        if Demographic_Category == "Overall":
            all_n_zero = False  # test
        all_n_zero = all(all(re.search("n = 0", x) for x in sample_sizes[1:]))
        if not all_n_zero:
            tabs = Crosstabs[Crosstabs.index("Prompt and Responses"):, :]

            questions = tabs[:, 0].unique()
            questions = [q for q in questions if q]

            subject = Group1

            document = Document(Template)

            for question in questions:
                tab = tabs[tabs[:, 0] == question]
                tab = tab[:5]

                stance_negative = tab[1, 1]
                stance_positive = tab[3, 1]

                text = Crosstabs[Crosstabs == question][1]

                lines = [f"{tools.toTitleCase(subject)} were asked to respond to the statement, \"{text}\"."]
                subject = subject.lower()

                data = []
                for group in legend:
                    index = legend.index(group)
                    group_original = group
                    if group == "Total":
                        group = subject
                    else:
                        group = f"those who selected \"{group}\""

                    group_results = tab[2:, index:index+3]
                    beg_mean = group_results[0, 0]
                    end_mean = group_results[0, 1]
                    change_mean = group_results[0, 2]
                    beg_support = convert_percentage(group_results[3, 0])
                    end_support = convert_percentage(group_results[3, 1])
                    change_support = convert_percentage(group_results[3, 2])
                    beg_opposition = convert_percentage(group_results[1, 0])
                    end_opposition = convert_percentage(group_results[1, 1])
                    change_opposition = convert_percentage(group_results[1, 2])

                    if beg_support == 0 or end_support == 0:
                        beg_support = convert_percentage(group_results[2, 0])
                        end_support = convert_percentage(group_results[2, 1])
                        change_support = convert_percentage(group_results[2, 2])
                        stance_positive = tab[2, 1]

                    if not pd.isnull(change_mean) and change_mean != "":
                        data.append((group_original, beg_support, end_support))

                        if extract_first_numeric(change_mean) != 0:
                            line1 = f"Among {group} ({sample_sizes[index]}), the mean {return_text(change_mean)} by {add_p_in_parenthesis(change_mean)}."
                        else:
                            line1 = f"Among {group}, the mean did not change."

                        if stance(end_support, end_opposition, stance_negative, stance_positive) == stance(beg_support, beg_opposition, stance_negative, stance_positive):
                            line2 = f"After deliberation, {stance_percentage(end_support, end_opposition, stance_negative, stance_positive)} among this group, similar to before deliberation."
                        else:
                            line2 = f"Before deliberation, {stance_percentage(beg_support, beg_opposition, stance_negative, stance_positive)} among this group, while after deliberation, {stance_percentage(end_support, end_opposition, stance_negative, stance_positive)}."

                        lines.append(f"{line1}. {line2}")

                if API_Key != "None":
                    def ask_chatgpt(prompt):
                        headers = {
                            "Authorization": f"Bearer {API_Key}",
                            "Content-Type": "application/json"
                        }
                        data = {
                            "model": "gpt-3.5-turbo",
                            "messages": [
                                {"role": "user", "content": prompt}
                            ]
                        }
                        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data)
                        response_json = response.json()
                        return response_json["choices"][0]["message"]["content"].strip()

                    lines = ask_chatgpt(f"Rephrase the following to sound less structured and more human, but do not change the numbers or conclusions. Keep the numbers the same. Don't remove the word deliberation: {lines}")

                if not lines:
                    lines = "API error. Ensure there is a valid API key for ChatGPT."

                if data:
                    df = pd.DataFrame(data, columns=["Group", f"Selected {stance_positive} Before Deliberation", f"Selected {stance_positive} After Deliberation"])
                    if Demographic_Category != "Overall":
                        df.set_index("Group", inplace=True)
                    else:
                        df = df.transpose()
                        df.columns = df.iloc[0]
                        df = df.iloc[1:]

                    # Create a table in the document
                    table = document.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                    table.autofit = False

                    # Set table style
                    table.style = "Table Grid"

                    # Set column widths
                    widths = [1, 2.5, 2.5]
                    for col, width in zip(table.columns, widths):
                        col.width = Pt(width * 72)

                    # Add table headers
                    headers = df.columns.tolist()
                    for i, header in enumerate(headers):
                        cell = table.cell(0, i)
                        cell.text = header
                        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # Add table data
                    for row, row_data in enumerate(dataframe_to_rows(df, index=True, header=False)):
                        for col, cell_value in enumerate(row_data):
                            cell = table.cell(row+1, col)
                            cell.text = str(cell_value)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                    # Set cell alignment
                    for row in table.rows:
                        for cell in row.cells:
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                question_paragraph = document.add_paragraph(question)
                question_paragraph.style = "Heading 1"

                for line in lines:
                    document.add_paragraph(line)

                document.add_paragraph()
                document.add_paragraph()

            file_name = File_Name.replace("Tables", "Report") + ".docx"
            file_path = Outputs + "/" + file_name
            document.save(file_path)
            print(f"Exported: {file_name}")

def write_excel(sample, name):

    name += ".xlsx"
    title = sample.name
    
    if len(title) > 31:
        title = title[:28] + "..."
    
    sample.crosstabs.index = [pd.NA] * len(sample.crosstabs)

    sample.crosstabs.to_excel(name, sheet_name = title, index = True, header = True)

    print(f"Exported: {name}")

def write_word(Crosstabs, Outputs, File_Name, Name_Group, Template, Document_Title, Type, Demographic_Category, Codebook):

    if Crosstabs.shape[1] < 14:
        File_Name = File_Name + ".docx"
        ColumnNumbers = Crosstabs.shape[1]
        RowNumbers = Crosstabs.shape[0]

        if Type == "Ordinal":
            # Creates a legend.
            Legend = Codebook.iloc[3:, Codebook.columns.get_loc(Demographic_Category)]
            Legend = Legend.dropna()
            Legend = pd.DataFrame(Legend)
            Legend = pd.concat([pd.DataFrame([["T", "Total"]]), Legend], ignore_index=True)

            # Gets the column names of the crosstabs
            ColumnNames = list(Crosstabs.iloc[4])

            # Gets parts of the crosstab to print
            Titles = Crosstabs.iloc[0]
            SampleSizes = Crosstabs.iloc[2]
            Crosstabs = Crosstabs.iloc[4:]

            # Converts the dataframes to tables
            Legend = Legend.to_html(index=False, header=False)
            Titles = Titles.to_frame().T.to_html(index=False, header=False)
            SampleSizes = SampleSizes.to_frame().T.to_html(index=False, header=False)
            Crosstabs = Crosstabs.to_html(index=False, header=False)

            # Creates the Word document for export
            document = Document(Template)
            document.add_paragraph()
            document.add_paragraph(Legend)
            document.add_paragraph()
            document.add_paragraph(Titles)
            document.add_paragraph()
            document.add_paragraph(SampleSizes)
            document.add_paragraph()
            document.add_paragraph(Crosstabs)
            document.save(Outputs + "/" + File_Name)

            # Notifies of document export
            print("Exported:", File_Name)

        elif Type == "Nominal":
            # Gets the column names of the crosstabs
            ColumnNames = list(Crosstabs.iloc[3])

            # Gets parts of the crosstab to print
            SampleSizes = Crosstabs.iloc[1]
            Crosstabs = Crosstabs.iloc[4:]

            # Converts the dataframes to tables
            SampleSizes = SampleSizes.to_frame().T.to_html(index=False, header=False)
            Crosstabs = Crosstabs.to_html(index=False, header=False)

            # Creates the Word document for export
            document = Document(Template)
            document.add_paragraph()
            document.add_paragraph(SampleSizes)
            document.add_paragraph()
            document.add_paragraph(Crosstabs)
            document.save(Outputs + "/" + File_Name)

            # Notifies of document export
            print("Exported:", File_Name)

    else:
        print("Cannot export Word version due to large file size or template error.")

def create_crosstab(type, data, index, columns, weight):

    if type == "Nominal":
        margins = False
        dropna = True
        normalize = False
    else:
        margins = True
        dropna = False
        normalize = 'columns'
    
    absolute_frequencies = pd.crosstab(
        index = data[index].cat.add_categories(['DK/NA']).fillna('DK/NA'),
        columns = data[columns],
        values = data[weight],
        aggfunc = 'sum',
        margins = margins,
        dropna = dropna,
        normalize = normalize)
    
    combined_frequencies = (absolute_frequencies / absolute_frequencies.sum().sum() * 100).round(1).apply(lambda x: x).applymap(lambda x: f"({x}%)") + ' ' + absolute_frequencies.astype(int).astype(str)

    if type == "Nominal":
        return combined_frequencies.iloc[:, ::-1]
    else:
        return 100*absolute_frequencies.iloc[:, ::-1]

def add_crosstab_tests(sample, nominal_variable, ordinal_variable):

        sample.means = pd.DataFrame([[pd.NA] * len(sample.crosstab.columns)], columns = sample.crosstab.columns)

        for filter in sample.one.crosstab.columns:
            
            mean1, mean2, mean_difference = test_t(sample, filter, nominal_variable, ordinal_variable)
            
            crosstab_index = list(sample.one.crosstab.columns).index(filter)
            
            sample.means.iloc[0, crosstab_index + 0*len(sample.one.crosstab.columns)] = mean1
            sample.means.iloc[0, crosstab_index + 1*len(sample.one.crosstab.columns)] = mean2
            sample.means.iloc[0, crosstab_index + 2*len(sample.one.crosstab.columns)] = mean_difference
            
            filter_samplesize = add_sample_size(filter, sample.one.values if filter == "All" else sample.one.values[sample.one.labels[nominal_variable] == filter])
            crosstab_header = sample.crosstab.columns.tolist()
            crosstab_header[crosstab_index + 0*len(sample.one.crosstab.columns)] = filter_samplesize
            crosstab_header[crosstab_index + 1*len(sample.one.crosstab.columns)] = filter_samplesize
            sample.crosstab.columns = sample.means.columns = pd.Index(crosstab_header)
        
        return combine_crosstabs(sample.means, sample.crosstab)

def combine_crosstabs(crosstab1, crosstab2):
    
    if len(crosstab1.columns) != len(crosstab2.columns):
        crosstab2 = pd.DataFrame(columns = crosstab1.columns)
    
    crosstab1.columns = crosstab2.columns = pd.MultiIndex.from_product([['Level1'], crosstab1.columns])

    crosstabs = pd.concat([crosstab1, crosstab2])

    crosstabs.columns = crosstabs.columns.get_level_values(1)

    return crosstabs

def test_chi(variable, observed, expected):

    observed_expected = np.column_stack((observed, expected))

    observed_expected = observed_expected[~np.apply_along_axis(lambda y: np.all(y == 0), 1, observed_expected)]

    _, P, _, _ = chi2_contingency(observed_expected)

    if not np.isnan(P):
        
        if np.any(expected < 5):
            return f"{variable} (P = {P:.3f}) Warning: P-value may be incorrect because at least one expected value is less than 5."
        else:
            return f"{variable} (P = {P:.3f})"
    
    return variable

def test_t(sample, filter, nominal_variable, ordinal_variable):
    
    sample.one.filtered = sample.one.values
    sample.two.filtered = sample.two.values

    if not filter == "All":
        sample.one.filtered = sample.one.values[sample.one.labels[nominal_variable] == filter]
        sample.two.filtered = sample.two.values[sample.two.labels[nominal_variable] == filter]
    
    if sample.paired:
        complete_cases = pd.concat((sample.one.filtered[ordinal_variable], sample.one.filtered[sample.one.weight], sample.two.filtered[ordinal_variable], sample.two.filtered[sample.two.weight]), axis = 1).dropna()
        sample.one.ordinal_filtered = complete_cases.groupby(level=0, axis=1).nth(0)[ordinal_variable]
        sample.one.weights_filtered = complete_cases.groupby(level=0, axis=1).nth(0)[sample.one.weight]
        sample.two.ordinal_filtered = complete_cases.groupby(level=0, axis=1).nth(1)[ordinal_variable]
        sample.two.weights_filtered = complete_cases.groupby(level=0, axis=1).nth(1)[sample.two.weight]
        
    else:
        complete_cases = pd.concat((sample.one.filtered[ordinal_variable], sample.one.filtered[sample.one.weight]), axis = 1).dropna()
        sample.one.ordinal_filtered = complete_cases[ordinal_variable]
        sample.one.weights_filtered = complete_cases[sample.one.weight]

        complete_cases = pd.concat((sample.two.filtered[ordinal_variable], sample.two.filtered[sample.two.weight]), axis = 1).dropna()
        sample.two.ordinal_filtered = complete_cases[ordinal_variable]
        sample.two.weights_filtered = complete_cases[sample.two.weight]

    if sample.paired:
        P = statsmodels.stats.weightstats.DescrStatsW(
            data = sample.two.ordinal_filtered - sample.one.ordinal_filtered,
            weights = sample.one.weights_filtered).ttest_mean(0)[1]

    else:
        P = statsmodels.stats.weightstats.ttest_ind(
            x1 = sample.one.ordinal_filtered,
            x2 = sample.two.ordinal_filtered,
            alternative = 'two-sided',
            usevar = 'unequal',
            weights = (sample.one.weights_filtered, sample.two.weights_filtered))[1]
    
    if not sum(sample.one.weights_filtered) == 0 and not sum(sample.two.weights_filtered) == 0:
        mean1 = np.average(sample.one.ordinal_filtered, weights = sample.one.weights_filtered)
        mean2 = np.average(sample.two.ordinal_filtered, weights = sample.two.weights_filtered)
        mean_difference = "{:.3f}".format(mean2 - mean1)     
        mean1 = "{:.3f}".format(mean1)
        mean2 = "{:.3f}".format(mean2)
    else:
        mean1 = mean2 = mean_difference = pd.NA
    
    if not np.isnan(P):
        mean_difference = f"{mean_difference} (P = {P:.3f})"
    
    return mean1, mean2, mean_difference

def document_title(sample, kind, type, nominal_variable):

        if sample.one.weight == sample.two.weight:
           weights = sample.one.weight
        else:
            weights = sample.one.weight + sample.two.weight
        
        if type == "Nominal":
            return " - ".join([kind, type, sample.name, weights])
        else:
            return " - ".join([kind, type, sample.name, weights, nominal_variable])

def comparison_name(sample1, sample2):

    if sample1.group == sample2.group:
        return f"{sample1.group} at {sample1.time} v. {sample2.time}"
    elif sample1.time == sample2.time:
        return f"{sample1.group} v. {sample2.group} at {sample1.time}"
    else:
        return f"{sample1.group} at {sample1.time} v. {sample2.group} at {sample2.time}"

def add_sample_size(variable, sample):

    return f"{variable} (n = {len(sample)})"

class subsample:
    def __init__(self, group, time, weight, values, labels):
        self.group = group
        self.time = time
        self.weight = weight
        self.name = f"{group} at {time}"
        self.values = values[(values['Group'] == group) & (values['Time'] == time)].assign(Overall = 1)
        self.labels = labels[(values['Group'] == group) & (values['Time'] == time)].assign(Overall = "Total")

analysis("Python/Dataset (Short).sav")