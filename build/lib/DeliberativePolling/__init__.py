import os
import re
import time
import threading
import pandas as pd
import numpy as np
import openpyxl
import warnings
import pyreadstat
import statsmodels
import statsmodels.stats.weightstats
from scipy.stats import chi2_contingency
from itertools import combinations, product
from tqdm import tqdm
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

warnings.filterwarnings("ignore")


def outputs(file, fast=False):
    """
    This function takes an .sav file from IBM SPSS Statistics and creates tables and reports.
    """

    assert file.lower().endswith(
        ".sav"
    ), 'File must be a .sav file from IBM SPSS Statistics. See pypi.org/project/DeliberativePolling for "How To" guide for this package.'

    values, codebook = pyreadstat.read_sav(file, apply_value_formats=False)
    labels = pyreadstat.read_sav(file, apply_value_formats=True)[0]

    scale_variables = [
        key for key, measure in codebook.variable_measure.items() if measure == "scale"
    ]
    weights = ["Unweighted"] + [
        var for var in scale_variables if "weight" in var.lower()
    ]

    check_numeric(codebook, weights[1:])

    for variable in ["Time", "Group", "ID"]:
        if variable not in values:
            raise ValueError(
                f'"{variable}" variable not found. See pypi.org/project/DeliberativePolling for "How To" guide for this package.'
            )

        if values[variable].isna().any():
            raise ValueError(f'Empty cells in "{variable}" variable found.')

    values["Time"] = labels["Time"]
    values["Group"] = labels["Group"]

    sample_comparisons = [
        comb
        for comb in list(
            combinations(
                list(
                    product(values["Group"].unique(), values["Time"].unique(), weights)
                ),
                2,
            )
        )
        if not (comb[0][0] == comb[1][0] and comb[0][2] != comb[1][2])
    ]

    for combination in tqdm(
        sample_comparisons,
        position=0,
        desc="Comparing Weighted Sampling",
        leave=True,
    ):

        class sample:
            one = subsample(
                combination[0][0],
                combination[0][1],
                combination[0][2],
                values,
                labels,
            )
            two = subsample(
                combination[1][0],
                combination[1][1],
                combination[1][2],
                values,
                labels,
            )
            all_values = values
            all_labels = labels
            metadata = codebook
            name = comparison_name(one, two)
            paired = str(one.group) + one.weight == str(two.group) + two.weight

        if sample.one.values["ID"].duplicated().any():
            raise ValueError(f'Duplicate IDs in "{sample.one.name}" found.')
        if sample.two.values["ID"].duplicated().any():
            raise ValueError(f'Duplicate IDs in "{sample.two.name}" found.')
        if (
            sample.paired
            and len(
                set(sample.one.values["ID"]).intersection(set(sample.two.values["ID"]))
            )
            == 0
        ):
            raise ValueError(
                f'No shared IDs between "{sample.one.name}" and "{sample.two.name}" were found. If the samples are not the same experimental group, they should not share the group name "{sample.one.group}".'
            )

        sample.one.values.set_index(sample.one.values["ID"], inplace=True)
        sample.two.values.set_index(sample.two.values["ID"], inplace=True)
        sample.one.labels.set_index(sample.one.values["ID"], inplace=True)
        sample.two.labels.set_index(sample.two.values["ID"], inplace=True)

        compare_samples(sample, "Nominal", fast)
        compare_samples(sample, "Ordinal", fast)

    print('Analysis complete. See "Outputs" folder in directory.')

    if "limit" in globals():
        print(limit)


def compare_samples(sample, type, fast):
    sample.crosstabs = pd.DataFrame()
    sample.summaries = pd.DataFrame(columns=["Variable", "Summary"])

    sample.metadata.variable_measure.pop("Group", None)
    sample.metadata.variable_measure.pop("Time", None)
    sample.metadata.variable_measure.pop("ID", None)
    nominal_variables = [
        key
        for key, measure in sample.metadata.variable_measure.items()
        if measure == "nominal"
    ]
    check_labels(sample, nominal_variables)

    if type == "Nominal":
        ordinal_variables = [1]
    else:
        ordinal_variables = list(
            reversed(
                [
                    key
                    for key, measure in sample.metadata.variable_measure.items()
                    if measure == "ordinal"
                ]
            )
        )
        check_labels(sample, ordinal_variables)

    if len(nominal_variables) == 0:
        raise ValueError(f"Nominal variables not found.")
    if len(ordinal_variables) == 0:
        raise ValueError(f"Ordinal variables not found.")

    for nominal_variable in tqdm(
        nominal_variables,
        position=1,
        desc="Comparing Nominal Variables",
        leave=False,
    ):
        shared_IDs = sample.one.values.index.intersection(sample.two.values.index)
        if (
            sample.paired
            and any(
                sample.one.values[nominal_variable].loc[shared_IDs]
                != sample.two.values[nominal_variable].loc[shared_IDs]
            )
            and type == "Ordinal"
        ):
            variations = [f" ({sample.one.time})", f" ({sample.two.time})"]

            sample.one.values.original_nominal = sample.one.values[
                nominal_variable
            ].loc[shared_IDs]
            sample.two.values.original_nominal = sample.two.values[
                nominal_variable
            ].loc[shared_IDs]
            sample.one.labels.original_nominal = sample.one.labels[
                nominal_variable
            ].loc[shared_IDs]
            sample.two.labels.original_nominal = sample.two.labels[
                nominal_variable
            ].loc[shared_IDs]

        else:
            variations = [""]

        for variation in variations:
            if variations.index(variation) == 0 and variation != "":
                sample.one.values[nominal_variable].loc[
                    shared_IDs
                ] = sample.one.values.original_nominal
                sample.two.values[nominal_variable].loc[
                    shared_IDs
                ] = sample.one.values.original_nominal
                sample.one.labels[nominal_variable].loc[
                    shared_IDs
                ] = sample.one.labels.original_nominal
                sample.two.labels[nominal_variable].loc[
                    shared_IDs
                ] = sample.one.labels.original_nominal

            if variations.index(variation) == 1:
                sample.one.values[nominal_variable].loc[
                    shared_IDs
                ] = sample.two.values.original_nominal
                sample.two.values[nominal_variable].loc[
                    shared_IDs
                ] = sample.two.values.original_nominal
                sample.one.labels[nominal_variable].loc[
                    shared_IDs
                ] = sample.two.labels.original_nominal
                sample.two.labels[nominal_variable].loc[
                    shared_IDs
                ] = sample.two.labels.original_nominal

            for ordinal_variable in tqdm(
                ordinal_variables,
                position=2,
                desc="Comparing Ordinal Variables",
                leave=False,
                total=len(ordinal_variables),
            ):
                if type == "Nominal":
                    sample.crosstab = nominal_crosstab(sample, nominal_variable)
                    sample.summary = nominal_summary(sample, nominal_variable)

                if type == "Ordinal":
                    sample.crosstab = ordinal_crosstab(
                        sample, nominal_variable, ordinal_variable
                    )
                    sample.summary = ordinal_summary(sample, ordinal_variable)

                sample.summary.columns = sample.summaries.columns
                sample.summaries = pd.concat([sample.summary, sample.summaries], axis=0)
                sample.crosstabs = crosstab_concat(sample.crosstab, sample.crosstabs)

            if type == "Ordinal":
                second_header = [""] * len(sample.crosstabs.columns)
                second_header[2] = sample.one.name
                second_header[
                    int((len(sample.crosstabs.columns) - 2) / 3) + 2
                ] = sample.two.name
                second_header[
                    int((len(sample.crosstabs.columns) - 2) / 3) * 2 + 2
                ] = "Difference"
                multi_col = pd.MultiIndex.from_tuples(
                    [
                        (second_header[i], col)
                        for i, col in enumerate(sample.crosstabs.columns)
                    ]
                )
                sample.crosstabs.columns = multi_col

                variable = sample.metadata.column_labels[
                    sample.metadata.column_names.index(nominal_variable)
                ]
                name = document_title(sample, type, variable) + variation

                write_xlsx(sample, name, variable)
                if not fast:
                    write_docx(sample, name, variable)
                sample.crosstabs = pd.DataFrame()
                sample.summaries = pd.DataFrame(columns=["Variable", "Summary"])

    if type == "Nominal":
        name = document_title(sample, type, nominal_variable)

        write_xlsx(sample, name)
        if not fast:
            write_docx(sample, name)


def nominal_crosstab(sample, nominal_variable):
    sample.one.crosstab = crosstab_create(
        type="Nominal",
        data=sample.one.labels,
        index=nominal_variable,
        columns="Total",
        weight=sample.one.weight,
    )

    sample.two.crosstab = crosstab_create(
        type="Nominal",
        data=sample.two.labels,
        index=nominal_variable,
        columns="Total",
        weight=sample.two.weight,
    )

    sample.crosstab = pd.concat([sample.one.crosstab, sample.two.crosstab], axis=1)

    sample.crosstab = sample.crosstab.reset_index()
    sample.crosstab.insert(0, "Variable", np.nan)
    sample.crosstab.columns = [
        "Variable",
        "Label",
        add_sample_size(sample.one.name, sample.one.values[nominal_variable]),
        add_sample_size(sample.two.name, sample.two.values[nominal_variable]),
    ]

    sample.crosstab.loc[0, "Variable"] = x_test(
        variable=sample.metadata.column_labels[
            sample.metadata.column_names.index(nominal_variable)
        ],
        observed=pd.crosstab(
            index=sample.one.labels[nominal_variable],
            columns="Total",
            values=sample.one.labels[sample.one.weight],
            aggfunc="sum",
        ),
        expected=pd.crosstab(
            index=sample.two.labels[nominal_variable],
            columns="Total",
            values=sample.two.labels[sample.one.weight],
            aggfunc="sum",
        ),
    )

    return sample.crosstab


def ordinal_crosstab(sample, nominal_variable, ordinal_variable):
    labels = sample.metadata.variable_value_labels[ordinal_variable].values()
    labels_ordered = []
    [labels_ordered.append(value) for value in labels if value not in labels_ordered]
    labels_ordered = labels_ordered + ["DK/NA"]

    sample.one.crosstab = crosstab_create(
        type="Ordinal",
        data=sample.one.labels,
        index=ordinal_variable,
        columns=nominal_variable,
        weight=sample.one.weight,
        labels=labels_ordered,
    )

    sample.two.crosstab = crosstab_create(
        type="Ordinal",
        data=sample.two.labels,
        index=ordinal_variable,
        columns=nominal_variable,
        weight=sample.two.weight,
        labels=labels_ordered,
    )

    sample.crosstab = pd.concat(
        [
            sample.one.crosstab,
            sample.two.crosstab,
            sample.two.crosstab - sample.one.crosstab,
        ],
        axis=1,
    )

    sample.crosstab = (
        sample.crosstab.fillna(0)
        .round(1)
        .apply(lambda x: x)
        .applymap(lambda x: f"{x}%")
    )

    sample.crosstab = crosstab_means(sample, nominal_variable, ordinal_variable)

    sample.crosstab = sample.crosstab.reset_index()
    sample.crosstab.insert(0, "Variable", np.nan)
    crosstab_header = sample.crosstab.columns.tolist()
    crosstab_header[0] = "Variable"
    crosstab_header[1] = "Label"
    sample.crosstab.columns = crosstab_header
    sample.crosstab.iloc[0, 0] = ordinal_variable
    sample.crosstab.iloc[0, 1] = sample.metadata.column_labels[
        sample.metadata.column_names.index(ordinal_variable)
    ]

    return sample.crosstab


def nominal_summary(sample, nominal_variable):
    statement = ""
    crosstab = sample.crosstab.copy()
    crosstab.set_index("Label", inplace=True)

    label = sample.metadata.column_labels[
        sample.metadata.column_names.index(nominal_variable)
    ]

    if "P = " in crosstab.iloc[0, 0]:
        difference = dist_comparison(crosstab.iloc[0, 0].split("P = ")[1].split(")")[0])

        likeness, plurality1, plurality2 = plurality_comparison(
            crosstab.iloc[:, 1],
            crosstab.iloc[:, 2],
        )

        statement = f'The distribution of "{label}" {difference} between {crosstab.columns[1]} and {crosstab.columns[2]}. In {sample.two.name}, there was {plurality2} among this group, {likeness} in {sample.one.name} {plurality1}.'

        if "Warning:" in crosstab.iloc[0, 0]:
            statement += " Note that the significance test may be incorrect because at least one expected value is less than 5."

    if len(statement) == 0:
        statement = "Insufficient data to generate prose summary."
    return pd.DataFrame([label, statement]).T


def ordinal_summary(sample, ordinal_variable):
    statement = ""
    crosstab = sample.crosstab.copy()
    crosstab.set_index("Label", inplace=True)

    label = sample.metadata.column_labels[
        sample.metadata.column_names.index(ordinal_variable)
    ]

    length = int((len(crosstab.columns) - 1) / 3)

    for value in crosstab.columns[1 : 1 + length]:
        index_one = list(crosstab.columns).index(value)
        index_two = index_one + length
        index_dif = index_two + length

        value2 = crosstab.columns[index_two]

        if type(crosstab.iloc[0, index_dif]) == str:
            likeness, plurality1, plurality2 = plurality_comparison(
                crosstab.iloc[1:-1, index_one],
                crosstab.iloc[1:-1, index_two],
            )

            difference = mean_comparison(sample, crosstab.iloc[0, index_dif])

            if sample.paired:
                if value.split(" (")[0] == "All":
                    statement += f"{sample.one.group} ({value.split(' (')[1]} responded to the statement, \"{label}\". The"
                else:
                    statement += f" Among those who selected \"{value.split(' (')[0]}\" ({value.split(' (')[1]}, the"
                statement += f" average response {difference} between {sample.one.time} and {sample.two.time}. At {sample.two.time}, there was {plurality2} among this group, {likeness} at {sample.one.time} {plurality1}"
                statement += f"."
            else:
                if value.split(" (")[0] == "All":
                    statement += f'{sample.one.name} and {sample.two.name} responded to the statement, "{label}". There'
                else:
                    statement += (
                        f" Among those who selected \"{value.split(' (')[0]}\", there"
                    )
                statement += f" {difference} in the average response between {sample.one.name} ({value.split(' (')[1]} and {sample.two.name} ({value2.split(' (')[1]}. Among {sample.two.group}, there was {plurality2}, {likeness} {sample.one.group} {plurality1}"
                statement += f"."
    if len(statement) == 0:
        statement = "Insufficient data to generate prose summary."
    return pd.DataFrame([ordinal_variable, statement]).T


def write_xlsx(sample, name, variable=""):
    name += ".xlsx"
    title = sample.name

    sheet_name = title
    if len(sheet_name) > 31:
        sheet_name = sheet_name[:28] + "..."

    sample.crosstabs.index = [pd.NA] * len(sample.crosstabs)

    if "Ordinal" in name:
        index = True
    else:
        index = False

    title = re.sub(r'[<>:"/\\|?*:]', "", title)
    name = re.sub(r'[<>:"/\\|?*:]', "", name)
    sheet_name = re.sub(r'[<>:"/\\|?*:]', "", sheet_name)

    directory = f"Outputs/{sample.one.time} v. {sample.two.time}/{title}"
    if os.path.exists(f"Outputs/{sample.two.time} v. {sample.one.time}"):
        directory = f"Outputs/{sample.two.time} v. {sample.one.time}/{title}"
    directory = f"{directory}/{variable}"
    os.makedirs(directory, exist_ok=True)

    sample.crosstabs.to_excel(
        f"{directory}/Tables - {name}",
        sheet_name=sheet_name,
        index=index,
        header=True,
    )

    sample.summaries.to_excel(
        f"{directory}/Report - {name}",
        sheet_name=sheet_name,
        index=False,
        header=True,
    )


def write_docx(sample, name, variable=""):
    try:
        sample.crosstabs.fillna("", inplace=True)
    except TypeError:
        pass

    name += ".docx"
    title = sample.name

    if variable == "":
        header_text = title
    else:
        header_text = f"{title} by {variable}"

    for sheet in [sample.summaries, sample.crosstabs]:
        document = Document()

        section = document.sections[0]

        if not len(sheet.columns) == 2:
            section.page_width = Inches(22)
            section.page_height = Inches(22)
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        header = document.add_heading(header_text, 0)
        for run in header.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

        sheets = []
        num_columns = len(sheet.columns)
        max_rows_per_split = 1
        start_idx = 0
        while start_idx < len(sheet):
            end_idx = start_idx + max_rows_per_split
            sheets.append(sheet.iloc[start_idx:end_idx])
            start_idx = end_idx

        iterations = -1
        for sheet in sheets:
            iterations += 1

            if iterations == 0 and "Ordinal" in name:
                spacer = 2
                multiplier = 2
            else:
                spacer = 0
                multiplier = 2

            if iterations == 0 and "Nominal" in name:
                spacer = 1
                multiplier = 1
            else:
                spacer = 0
                multiplier = 2

            rows, cols = sheet.shape
            table = document.add_table(rows=rows + multiplier * spacer, cols=cols)
            table.style = "Medium List 2"

            if iterations == 0:
                for i, column in enumerate(sheet.columns):
                    cell = table.cell(0, i)
                    for p in cell.paragraphs:
                        p.clear()
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    if isinstance(column, tuple):
                        p1 = cell.add_paragraph(str(column[0]))
                        set_font(p1)
                        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        split = str(column[1]).find("(n = ")
                        if split != -1:
                            parts = [
                                str(column[1])[: split - 1],
                                str(column[1])[split:],
                            ]
                            for part in parts:
                                p = cell.add_paragraph(part)
                                set_font(p)
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            p2 = cell.add_paragraph(str(column[1]))
                            set_font(p2)
                            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        p = cell.add_paragraph(str(column))
                        set_font(p)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    vertical_alignment(cell)

            for i, row in enumerate(sheet.iterrows()):
                data = row[1]
                for j, value in enumerate(data):
                    cell = table.cell(i + spacer, j)
                    cell.text = str(value)
                    set_font(cell)

                    for paragraph in cell.paragraphs:
                        if j < 2:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        else:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    if not len(sheet.columns) == 2:
                        vertical_alignment(cell)

        title = re.sub(r'[<>:"/\\|?*:]', "", title)
        name = re.sub(r'[<>:"/\\|?*:]', "", name)

        directory = f"Outputs/{sample.one.time} v. {sample.two.time}/{title}"
        if os.path.exists(f"Outputs/{sample.two.time} v. {sample.one.time}"):
            directory = f"Outputs/{sample.two.time} v. {sample.one.time}/{title}"
        directory = f"{directory}/{variable}"
        os.makedirs(directory, exist_ok=True)

        if len(sheet.columns) == 2:
            document.save(f"{directory}/Report - {name}")
        else:
            document.save(f"{directory}/Tables - {name}")


def crosstab_create(type, data, index, columns, weight, labels=None):
    if type == "Nominal":
        margins = False
        dropna = True
        normalize = False
        index_data = data[index]
    else:
        margins = True
        dropna = False
        normalize = "columns"
        index_data = data[index].cat.add_categories(["DK/NA"]).fillna("DK/NA")

    absolute_frequencies = pd.crosstab(
        index=index_data,
        columns=data[columns],
        values=data[weight],
        aggfunc="sum",
        margins=margins,
        dropna=dropna,
        normalize=normalize,
    )

    if type == "Nominal":
        combined_frequencies = (
            (absolute_frequencies / absolute_frequencies.sum().sum() * 100)
            .round(1)
            .astype(str)
            .replace("nan", "0.0")
            .applymap(lambda x: f"({x}%)")
            + " "
            + absolute_frequencies.round().astype(int).astype(str)
        )

        return combined_frequencies.iloc[:, ::-1].replace("nan", "0.0")
    else:
        return 100 * absolute_frequencies.iloc[:, ::-1].fillna(0).reindex(
            labels
        ).replace("nan", "0")


def crosstab_concat(crosstab1, crosstab2):
    if len(crosstab1.columns) != len(crosstab2.columns):
        crosstab2 = pd.DataFrame(columns=crosstab1.columns)

    crosstab1.columns = crosstab2.columns = pd.MultiIndex.from_product(
        [["Level1"], crosstab1.columns]
    )

    crosstabs = pd.concat([crosstab1, crosstab2])

    crosstabs.columns = crosstabs.columns.get_level_values(1)

    return crosstabs


def crosstab_means(sample, nominal_variable, ordinal_variable):
    sample.means = pd.DataFrame(
        [[pd.NA] * len(sample.crosstab.columns)], columns=sample.crosstab.columns
    )

    for filter in sample.one.crosstab.columns:
        mean1, mean2, mean_difference = t_test(
            sample, filter, nominal_variable, ordinal_variable
        )

        crosstab_index = list(sample.one.crosstab.columns).index(filter)

        sample.means.iloc[
            0, crosstab_index + 0 * len(sample.one.crosstab.columns)
        ] = mean1
        sample.means.iloc[
            0, crosstab_index + 1 * len(sample.one.crosstab.columns)
        ] = mean2
        sample.means.iloc[
            0, crosstab_index + 2 * len(sample.one.crosstab.columns)
        ] = mean_difference

        crosstab_header = sample.crosstab.columns.tolist()
        crosstab_header[
            crosstab_index + 0 * len(sample.one.crosstab.columns)
        ] = add_sample_size(
            filter,
            sample.one.values
            if filter == "All"
            else sample.one.values[sample.one.labels[nominal_variable] == filter],
        )
        crosstab_header[
            crosstab_index + 1 * len(sample.one.crosstab.columns)
        ] = add_sample_size(
            filter,
            sample.two.values
            if filter == "All"
            else sample.two.values[sample.two.labels[nominal_variable] == filter],
        )
        sample.crosstab.columns = sample.means.columns = pd.Index(crosstab_header)

    return crosstab_concat(sample.means, sample.crosstab)


def x_test(variable, observed, expected):
    combined_index = observed.index.union(expected.index)
    observed = observed.reindex(combined_index, fill_value=0)
    expected = expected.reindex(combined_index, fill_value=0)

    observed_expected = np.column_stack((observed, expected))

    observed_expected = observed_expected[
        ~np.apply_along_axis(lambda y: np.all(y == 0), 1, observed_expected)
    ]

    if not type(sum(observed_expected)) == int and not np.any(
        np.all(observed_expected == 0, axis=0)
    ):
        _, P, _, _ = chi2_contingency(observed_expected)
    else:
        P = np.nan

    if not np.isnan(P):
        if np.any(expected < 5):
            return f"{variable} (P = {P:.3f}) Warning: P-value may be incorrect because at least one expected value is less than 5."
        else:
            return f"{variable} (P = {P:.3f})"

    return variable


def t_test(sample, filter, nominal_variable, ordinal_variable):
    sample.one.filtered = sample.one.values
    sample.two.filtered = sample.two.values

    if not filter == "All":
        sample.one.filtered = sample.one.values[
            sample.one.labels[nominal_variable] == filter
        ]
        sample.two.filtered = sample.two.values[
            sample.two.labels[nominal_variable] == filter
        ]

    sample = full_entries(sample, ordinal_variable)

    if sample.paired and len(sample.one.ordinal_filtered) == len(
        sample.two.ordinal_filtered
    ):
        P = statsmodels.stats.weightstats.DescrStatsW(
            data=sample.two.ordinal_filtered - sample.one.ordinal_filtered,
            weights=sample.one.weights_filtered,
        ).ttest_mean(0)[1]

    else:
        P = statsmodels.stats.weightstats.ttest_ind(
            x1=sample.one.ordinal_filtered,
            x2=sample.two.ordinal_filtered,
            alternative="two-sided",
            usevar="unequal",
            weights=(sample.one.weights_filtered, sample.two.weights_filtered),
        )[1]

    if not sum(sample.one.weights_filtered) == 0:
        mean1 = np.average(
            sample.one.ordinal_filtered, weights=sample.one.weights_filtered
        )
    else:
        mean1 = pd.NA

    if not sum(sample.two.weights_filtered) == 0:
        mean2 = np.average(
            sample.two.ordinal_filtered, weights=sample.two.weights_filtered
        )
    else:
        mean2 = pd.NA

    if not pd.isna(mean1) and not pd.isna(mean2):
        mean_difference = "{:.3f}".format(mean2 - mean1)
    else:
        mean_difference = pd.NA

    if not pd.isna(mean1):
        mean1 = "{:.3f}".format(mean1)
    if not pd.isna(mean2):
        mean2 = "{:.3f}".format(mean2)

    if not np.isnan(P):
        mean_difference = f"{mean_difference} (P = {P:.3f})"

    return mean1, mean2, mean_difference


def full_entries(sample, ordinal_variable):
    if (
        sample.paired
        and not all(np.isnan(value) for value in sample.one.filtered[ordinal_variable])
        and not all(np.isnan(value) for value in sample.two.filtered[ordinal_variable])
    ):
        complete_cases = pd.concat(
            (
                sample.one.filtered[ordinal_variable],
                sample.one.filtered[sample.one.weight],
                sample.two.filtered[ordinal_variable],
                sample.two.filtered[sample.two.weight],
            ),
            axis=1,
        ).dropna()
        sample.one.ordinal_filtered = complete_cases.groupby(level=0, axis=1).nth(0)[
            ordinal_variable
        ]
        sample.one.weights_filtered = complete_cases.groupby(level=0, axis=1).nth(0)[
            sample.one.weight
        ]
        sample.two.ordinal_filtered = complete_cases.groupby(level=0, axis=1).nth(1)[
            ordinal_variable
        ]
        sample.two.weights_filtered = complete_cases.groupby(level=0, axis=1).nth(1)[
            sample.two.weight
        ]

    else:
        complete_cases = pd.concat(
            (
                sample.one.filtered[ordinal_variable],
                sample.one.filtered[sample.one.weight],
            ),
            axis=1,
        ).dropna()
        sample.one.ordinal_filtered = complete_cases[ordinal_variable]
        sample.one.weights_filtered = complete_cases[sample.one.weight]

        complete_cases = pd.concat(
            (
                sample.two.filtered[ordinal_variable],
                sample.two.filtered[sample.two.weight],
            ),
            axis=1,
        ).dropna()
        sample.two.ordinal_filtered = complete_cases[ordinal_variable]
        sample.two.weights_filtered = complete_cases[sample.two.weight]

    return sample


def document_title(sample, type, variable):
    if type == "Nominal":
        return " - ".join([f"{type} Variables", sample.name])
    else:
        return " - ".join([f"{type} Variables", sample.name, variable])


def comparison_name(sample1, sample2):
    if sample1.weight == sample2.weight:
        if sample1.group == sample2.group:
            return f"{sample1.group} at {sample1.time} v. {sample2.time} ({sample1.weight})"
        elif sample1.time == sample2.time:
            return f"{sample1.group} v. {sample2.group} at {sample1.time} ({sample1.weight})"
        else:
            return f"{sample1.group} at {sample1.time} v. {sample2.group} at {sample2.time} ({sample1.weight})"
    else:
        if sample1.group == sample2.group:
            return f"{sample1.group} at {sample1.time} ({sample1.weight}) v. {sample2.time} ({sample2.weight})"
        elif sample1.time == sample2.time:
            return f"{sample1.group} ({sample1.weight}) v. {sample2.group} ({sample2.weight}) at {sample1.time}"
        else:
            return f"{sample1.group} at {sample1.time} ({sample1.weight}) v. {sample2.group} at {sample2.time} ({sample2.weight})"


def add_sample_size(variable, sample):
    return f"{variable} (n = {len(sample)})"


def set_font(item):
    if hasattr(item, "paragraphs"):
        for paragraph in item.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
    else:
        for run in item.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def vertical_alignment(cell, align="bottom"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def mean_comparison(sample, difference):
    numeric = float(difference.split(" (")[0])
    if "P =" in difference:
        P = float(difference.split("P = ")[1][:-1])
    else:
        P = 1

    if numeric < 0:
        change = "decreased"
    elif numeric > 0:
        change = "increased"
    else:
        change = "did not change"

    if sample.paired:
        if P < 0.05:
            return f"{change} significantly by {difference}"
        else:
            return f"did not change significantly (P = {P})"
    else:
        if P < 0.05:
            return f"was a significant difference by {difference}"
        else:
            return f"was not a significant (P = {P}) difference"


def dist_comparison(P):
    if float(P) < 0.05:
        return f"was significantly (P = {P}) different"
    else:
        return f"was not significantly (P = {P}) different"


def plurality(percentages):
    if percentages.apply(lambda x: "%)" in x).any():
        percentages = percentages.apply(lambda x: x.split("%")[0][1:] + "%")

    value = percentages.str.rstrip("%").astype(float).idxmax()
    numeric = float(percentages[value].rstrip("%"))

    if numeric / 100 < 0.5:
        kind = "plurality"
    elif numeric / 100 > 2 / 3:
        kind = "supermajority"
    else:
        kind = "majority"

    return value, numeric, kind


def plurality_comparison(percentages1, percentages2):
    value1, numeric1, kind1 = plurality(percentages1)
    value2, numeric2, kind2 = plurality(percentages2)

    if kind1 == kind2:
        likeness = "like"
        plurality1 = f"({numeric1}%)"
    else:
        likeness = "unlike"
        plurality1 = f'which had a {kind1} response of "{value1}" ({numeric1}%)'

    plurality2 = f'a {kind2} response of "{value2}" ({numeric2}%)'

    return likeness, plurality1, plurality2


def check_labels(sample, variables):
    for variable in variables:
        if not variable in sample.metadata.variable_value_labels:
            raise ValueError(f'Value labels for variable "{variable}" not found.')

        missing_categories = [
            category
            for category in sample.all_labels[variable].cat.categories
            if category not in sample.metadata.variable_value_labels[variable].values()
        ]
        if len(missing_categories) > 0:
            raise ValueError(
                f'Value labels for {missing_categories} in variable "{variable}" not found.'
            )

        if type(
            sample.metadata.column_labels[sample.metadata.column_names.index(variable)]
        ) is type(None):
            raise ValueError(f'Column label for variable "{variable}" not found.')


def check_numeric(codebook, weights):
    ordinal_variables_non_numeric = list(
        filter(
            lambda x: x not in ["Time", "Group", "ID"],
            {
                var: codebook.readstat_variable_types[var]
                for var, measure in codebook.variable_measure.items()
                if measure == "ordinal"
                and codebook.readstat_variable_types[var] != "double"
            },
        )
    )
    if len(ordinal_variables_non_numeric) > 0:
        raise ValueError(
            f'Ordinal variables {list(ordinal_variables_non_numeric)} found with Type other than "Numeric". Ensure all ordinal variables are of Type "Numeric".'
        )

    nominal_variables_non_numeric = list(
        filter(
            lambda x: x not in ["Time", "Group", "ID"],
            {
                var: codebook.readstat_variable_types[var]
                for var, measure in codebook.variable_measure.items()
                if measure == "nominal"
                and codebook.readstat_variable_types[var] != "double"
            },
        )
    )
    if len(nominal_variables_non_numeric) > 0:
        raise ValueError(
            f'Nominal variables {list(nominal_variables_non_numeric)} found with Type other than "Numeric". Ensure all nominal variables are of Type "Numeric".'
        )

    weight_variables_non_numeric = {
        var: codebook.readstat_variable_types[var]
        for var in weights
        if codebook.readstat_variable_types[var] != "double"
    }
    list(
        filter(
            lambda x: x not in ["Time", "Group", "ID"],
            {
                var: codebook.readstat_variable_types[var]
                for var in weights
                if codebook.variable_measure.items()
                if codebook.readstat_variable_types[var] != "double"
            },
        )
    )
    if len(weight_variables_non_numeric) > 0:
        raise ValueError(
            f'Weight variables {list(weight_variables_non_numeric.keys())} found with Type other than "Numeric". Ensure all weight variables are of Type "Numeric".'
        )


class subsample:
    def __init__(self, group, time, weight, values, labels):
        self.group = group
        self.time = time
        self.weight = weight
        self.name = f"{group} at {time}"
        self.values = (
            values[(values["Group"] == group) & (values["Time"] == time)]
            .assign(Total="Total")
            .assign(Unweighted=1)
        )
        self.labels = (
            labels[(values["Group"] == group) & (values["Time"] == time)]
            .assign(Total="Total")
            .assign(Unweighted=1)
        )
